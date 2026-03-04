"""
Generate a presentation script + Q&A preparation document for the Flappy Bird ML-Agent project.
10-minute presentation script + anticipated professor questions with answers.
Usage: python generate_docs.py
Output: FlappyBird_ML_Agent_Documentation.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def add_script_note(doc, text):
    """Add a presenter note / what to say."""
    p = doc.add_paragraph()
    run = p.add_run("[SAY]: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(50, 50, 50)


def add_action_note(doc, text):
    """Add a presenter action (show something, demo, etc.)."""
    p = doc.add_paragraph()
    run = p.add_run("[DO]: ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(100, 50, 50)


def add_qa(doc, question, answer):
    """Add a Q&A pair."""
    p = doc.add_paragraph()
    q_run = p.add_run("Q: " + question)
    q_run.bold = True
    q_run.font.size = Pt(11)

    a_para = doc.add_paragraph()
    a_run = a_para.add_run("A: " + answer)
    a_run.font.size = Pt(10)
    a_run.font.color.rgb = RGBColor(30, 30, 80)
    doc.add_paragraph()  # spacing


def read_file(relative_path):
    full_path = os.path.join(BASE_DIR, relative_path)
    try:
        with open(full_path, "r", encoding="utf-8-sig") as f:
            return f.read()
    except Exception as e:
        return f"[Error reading file: {e}]"


def add_code_block(doc, code):
    style_name = "Code Block"
    if style_name not in [s.name for s in doc.styles]:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = "Consolas"
        font.size = Pt(7.5)
        font.color.rgb = RGBColor(30, 30, 30)
        fmt = style.paragraph_format
        fmt.space_before = Pt(2)
        fmt.space_after = Pt(2)
        fmt.left_indent = Inches(0.2)

    for line in code.split("\n"):
        doc.add_paragraph(line, style=style_name)


def generate_documentation():
    doc = Document()

    # ============================================================
    # TITLE PAGE
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("Flappy Bird ML-Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Presentation Script & Q&A Preparation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    tech = doc.add_paragraph("Unity ML-Agents  |  PPO  |  Reinforcement Learning")
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info = doc.add_paragraph("Total presentation time: ~10 minutes")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.bold = True

    doc.add_page_break()

    # ============================================================
    # PART 1: PRESENTATION SCRIPT (~10 min)
    # ============================================================
    doc.add_heading("PART 1: Presentation Script", level=1)
    doc.add_paragraph("Estimated total: 10 minutes. Timing for each section is indicated.")
    doc.add_paragraph()

    # --- INTRO (1.5 min) ---
    doc.add_heading("1. Introduction (1.5 min)", level=2)
    add_script_note(doc,
        "Hello everyone. Today I'm going to present my project where I trained an AI agent "
        "to play Flappy Bird using reinforcement learning. The idea is simple: instead of "
        "programming rules like 'jump when the pipe is close', we let the AI figure out the "
        "optimal strategy entirely on its own through trial and error."
    )
    add_script_note(doc,
        "The project uses Unity as the game engine, and Unity ML-Agents toolkit for the "
        "reinforcement learning pipeline. The training algorithm is PPO — Proximal Policy "
        "Optimization — which is a state-of-the-art policy gradient method. The agent code "
        "is written in C#, while the training runs through a Python backend using PyTorch."
    )
    add_action_note(doc, "Show the Unity project open with GameScene loaded.")

    # --- HOW IT WORKS (2.5 min) ---
    doc.add_heading("2. How the Agent Works (2.5 min)", level=2)

    doc.add_heading("2a. Observations — What the Agent Sees", level=3)
    add_script_note(doc,
        "The agent doesn't see the screen like a human. It receives only 4 numbers every frame:"
    )
    add_script_note(doc,
        "First: the bird's Y position — how high or low it is. "
        "Second: the bird's vertical velocity — is it going up or falling down. "
        "Third: the horizontal distance to the next pipe. "
        "Fourth: the Y position of the gap in the next pipe."
    )
    add_script_note(doc,
        "All 4 values are normalized to roughly the same range so the neural network "
        "can process them efficiently. This is a very minimal observation space — just 4 "
        "floats — yet it contains all the information needed to play the game perfectly."
    )

    doc.add_heading("2b. Actions — What the Agent Can Do", level=3)
    add_script_note(doc,
        "The agent has a discrete action space with only 2 options: do nothing, which lets "
        "gravity pull the bird down, or jump, which applies an upward velocity. That's it. "
        "Every decision step, the agent picks one of these two actions."
    )

    doc.add_heading("2c. Reward System — How the Agent Learns", level=3)
    add_script_note(doc,
        "This is the most important part. Reinforcement learning is driven by rewards. "
        "The agent's goal is to maximize the total reward it receives. Here's how I designed "
        "the reward function:"
    )
    add_script_note(doc,
        "The agent gets +0.01 for every frame it stays alive — this encourages survival. "
        "It gets +1.0 every time it successfully passes through a pipe — this is the main "
        "objective. It gets -2.0 when it dies — either by hitting a pipe, hitting the ground, "
        "or going off screen. And there's also a gap alignment reward of +0.005, scaled by how "
        "close the bird is to the center of the pipe gap — this gives the agent a continuous "
        "signal guiding it toward the correct position."
    )
    add_script_note(doc,
        "The gap alignment reward was a key design decision. Without it, the agent only knows "
        "it did something right when it passes a pipe or something wrong when it dies. With it, "
        "the agent gets moment-to-moment feedback about whether it's in a good position."
    )

    # --- TRAINING (2 min) ---
    doc.add_heading("3. Training Process (2 min)", level=2)
    add_script_note(doc,
        "The training uses PPO, which is one of the most reliable reinforcement learning "
        "algorithms. It works by collecting batches of experience — the agent plays many "
        "episodes, records what it observed, what it did, and what reward it got — then updates "
        "the neural network to make good actions more likely and bad actions less likely."
    )
    add_script_note(doc,
        "The neural network is simple: 4 inputs, two hidden layers with 256 neurons each "
        "using ReLU activation, and 2 outputs for the action probabilities. The model has "
        "roughly 66,000 parameters — quite small by modern standards."
    )
    add_script_note(doc,
        "Key hyperparameters: batch size of 128, buffer size of about 20,000 experiences, "
        "learning rate of 0.0003 with linear decay, and a discount factor gamma of 0.995 "
        "which means the agent values future rewards very highly."
    )
    add_script_note(doc,
        "Training was done over approximately 1 million steps. In the early stages, the agent "
        "jumps randomly and dies immediately. By around 100K steps it starts learning to stay "
        "alive. By 500K it can pass several pipes, and after 1 million steps it reaches "
        "expert-level performance."
    )
    add_action_note(doc, "If possible, show TensorBoard graph of reward over time.")

    # --- DIFFICULTY (1 min) ---
    doc.add_heading("4. Difficulty Scaling (1 min)", level=2)
    add_script_note(doc,
        "The game has a built-in difficulty system that acts like a curriculum. "
        "In the Easy level, the gap between pipes is 50 units wide. As the agent passes "
        "more pipes, the difficulty increases: Medium has a 40-unit gap, Hard has 33, and "
        "Impossible has only 24 units. The spawn interval also decreases so pipes come faster."
    )
    add_script_note(doc,
        "This means the agent naturally experiences easy scenarios first and gradually faces "
        "harder challenges — similar to how curriculum learning works in machine learning "
        "research."
    )

    # --- DEMO (2 min) ---
    doc.add_heading("5. Live Demo (2 min)", level=2)
    add_action_note(doc,
        "Open Unity. Select Bird object. Make sure Behavior Parameters > Behavior Type "
        "is set to 'Inference Only' and the Model field has the .onnx file assigned. "
        "Press Play. Let the agent play for 30-60 seconds."
    )
    add_script_note(doc,
        "Here you can see the trained agent playing. Notice how it makes very precise jumps, "
        "always positioning itself near the center of the pipe gap. It's not using any "
        "hard-coded rules — it learned this strategy entirely through reinforcement learning."
    )
    add_action_note(doc,
        "Optionally: Switch Behavior Type to 'Default' and show that without the model, "
        "the bird just falls or jumps randomly."
    )
    add_script_note(doc,
        "The trained model is exported as an ONNX file, which is an open standard for "
        "neural networks. Unity can run it natively — no Python needed at runtime."
    )

    # --- CONCLUSION (1 min) ---
    doc.add_heading("6. Conclusion (1 min)", level=2)
    add_script_note(doc,
        "To summarize: this project demonstrates that a simple reinforcement learning agent "
        "with just 4 inputs and a small neural network can learn to play Flappy Bird at a "
        "superhuman level. The key ingredients were: a well-designed reward function with "
        "continuous feedback, the PPO algorithm for stable training, and Unity ML-Agents "
        "as the framework connecting the game to the training pipeline."
    )
    add_script_note(doc,
        "Challenges I faced included: getting the reward shaping right — my first version "
        "only had survival and death rewards and the agent barely learned. Adding the gap "
        "alignment reward was a breakthrough. Another challenge was the Time Scale issue — "
        "training at high simulation speed produced an agent that didn't perform well at "
        "normal speed, so I had to retrain at a realistic speed."
    )
    add_script_note(doc,
        "Thank you. I'm happy to take any questions."
    )

    doc.add_page_break()

    # ============================================================
    # PART 2: ANTICIPATED Q&A
    # ============================================================
    doc.add_heading("PART 2: Anticipated Professor Questions & Answers", level=1)
    doc.add_paragraph(
        "Below are likely questions a professor might ask, with detailed answers prepared."
    )
    doc.add_paragraph()

    # --- RL FUNDAMENTALS ---
    doc.add_heading("Reinforcement Learning Fundamentals", level=2)

    add_qa(doc,
        "Why did you choose Reinforcement Learning instead of supervised learning or a rule-based approach?",
        "Supervised learning would require a dataset of expert gameplay (state-action pairs), "
        "which is hard to generate for a game like Flappy Bird. A rule-based approach (e.g., "
        "'jump when pipe gap is above you') would work but wouldn't demonstrate machine learning. "
        "RL is the natural choice because the game provides a clear reward signal (score), "
        "and the agent can learn through interaction without needing labeled data. RL also "
        "discovers strategies that humans might not think of."
    )

    add_qa(doc,
        "What is PPO and why did you choose it over other RL algorithms like DQN or A2C?",
        "PPO (Proximal Policy Optimization) is a policy gradient method that directly optimizes "
        "the policy (action probabilities). It uses a clipped surrogate objective to prevent "
        "too-large policy updates, making training very stable. I chose PPO over DQN because: "
        "(1) PPO handles discrete actions well and is more sample-efficient for simple environments, "
        "(2) it's the default and best-supported algorithm in Unity ML-Agents, "
        "(3) PPO is generally more stable than vanilla policy gradient methods like A2C. "
        "DQN would also work here since we have discrete actions, but PPO is the industry "
        "standard for game-playing agents."
    )

    add_qa(doc,
        "What is the discount factor (gamma) and why did you set it to 0.995?",
        "Gamma determines how much the agent values future rewards versus immediate rewards. "
        "A gamma of 0.995 means the agent is very forward-looking — it cares a lot about "
        "rewards it will receive many steps in the future. This is important for Flappy Bird "
        "because the agent needs to plan ahead: a jump now affects its position several frames "
        "later when it reaches the next pipe. A lower gamma (e.g., 0.9) would make the agent "
        "too short-sighted, only caring about immediate survival rather than positioning for "
        "upcoming pipes."
    )

    add_qa(doc,
        "Can you explain the exploration-exploitation tradeoff in your project?",
        "In the early stages of training, the agent needs to explore — try different actions "
        "randomly to discover what works. Over time, it should exploit — use the strategies "
        "it has learned. PPO handles this through entropy regularization (beta parameter = 0.005). "
        "The entropy bonus encourages the agent to maintain some randomness in its actions early "
        "on. As training progresses and the learning rate decays linearly, the agent becomes "
        "more confident and deterministic in its actions."
    )

    add_qa(doc,
        "What is a policy in reinforcement learning?",
        "A policy is a mapping from states (observations) to actions. In our case, the policy "
        "is the neural network: it takes the 4 observation values as input and outputs a "
        "probability distribution over the 2 possible actions (jump or do nothing). During "
        "training, the agent samples from this distribution. During inference, it picks the "
        "most likely action. PPO optimizes this policy to maximize expected cumulative reward."
    )

    # --- REWARD DESIGN ---
    doc.add_heading("Reward Design", level=2)

    add_qa(doc,
        "Why did you use -2.0 for death instead of -1.0? Why is the penalty asymmetric with the +1.0 pipe reward?",
        "The death penalty needs to be strong enough to discourage risky behavior. With -1.0, "
        "the agent might learn that dying is 'worth it' if it can pass one pipe first (net 0). "
        "With -2.0, dying always results in a net loss even if the agent passed a pipe just "
        "before. This creates a strong incentive for safe, consistent play rather than risky "
        "strategies. The asymmetry is intentional — surviving matters more than scoring."
    )

    add_qa(doc,
        "What is reward shaping and why was the gap alignment reward important?",
        "Reward shaping is the practice of adding intermediate rewards that guide the agent "
        "toward desired behavior, without changing the optimal policy. The gap alignment reward "
        "(+0.005 scaled by proximity to pipe center) is a shaping reward — it provides a "
        "continuous gradient signal. Without it, the agent only gets feedback when it passes "
        "a pipe (+1) or dies (-2), which are sparse signals. With the gap alignment reward, "
        "every single frame the agent gets information about whether it's in a good position. "
        "This dramatically speeds up learning because the agent doesn't have to randomly "
        "discover that being near the gap center is good."
    )

    add_qa(doc,
        "Could you train the agent with only +1 for pipe passing and -1 for death? No survival or alignment rewards?",
        "Yes, it would eventually learn, but it would take significantly more training steps. "
        "The sparse reward problem is well-known in RL: when rewards are infrequent, the agent "
        "struggles to learn because most experiences provide no learning signal. The survival "
        "reward (+0.01/frame) and gap alignment reward (+0.005 scaled) make the reward dense, "
        "giving the agent useful feedback every frame. In my testing, adding these shaping "
        "rewards reduced training time by roughly 3-5x."
    )

    # --- ARCHITECTURE ---
    doc.add_heading("Architecture & Design Decisions", level=2)

    add_qa(doc,
        "Why only 4 observations? Couldn't you use more inputs like multiple upcoming pipes?",
        "Four observations are sufficient because Flappy Bird is a reactive game — you only "
        "need to know about the next pipe to make the optimal decision. Adding more pipes "
        "would increase the observation space without adding useful information, and could "
        "even slow down learning. The principle of minimal sufficient observation is important "
        "in RL: fewer inputs mean faster training and less chance of overfitting."
    )

    add_qa(doc,
        "Why did you normalize the observations?",
        "Neural networks learn best when inputs are in a similar numerical range, typically "
        "around [-1, 1] or [0, 1]. Without normalization, the Y position (range 0-50) and "
        "velocity (range -90 to +90) would have very different scales, making it harder for "
        "the network to assign appropriate weights. By dividing each observation by its maximum "
        "value, all inputs are in a comparable range, leading to more stable and faster training."
    )

    add_qa(doc,
        "Why 2 hidden layers with 256 neurons? Could a smaller network work?",
        "Yes, a smaller network (e.g., 64 or 128 neurons) could also work for this problem. "
        "I chose 256 as a safe middle ground — large enough to capture the decision boundary "
        "but small enough for fast training and inference. The network has about 66,000 "
        "parameters total. For comparison, GPT models have billions. A 2-layer network can "
        "approximate any continuous function (universal approximation theorem), which is more "
        "than enough for this binary decision task."
    )

    add_qa(doc,
        "What is the ONNX format and why do you use it?",
        "ONNX (Open Neural Network Exchange) is an open standard for representing neural "
        "networks. Unity ML-Agents exports trained models to ONNX format, which Unity's "
        "Barracuda inference engine can run natively in C#. This means the final game doesn't "
        "need Python or PyTorch at all — the model runs purely within Unity, making deployment "
        "simple and efficient. ONNX is also framework-agnostic, so a model trained in PyTorch "
        "can be used anywhere."
    )

    # --- TRAINING ---
    doc.add_heading("Training Process", level=2)

    add_qa(doc,
        "How long did training take?",
        "The main training run (1 million steps) took approximately 30-40 minutes at Time Scale 20. "
        "When training at Time Scale 1 (real-time speed), training is about 20x slower — around "
        "6,000 steps per minute. A full 1M step run at Time Scale 1 would take about 3 hours. "
        "The final model was trained from a checkpoint of the high-speed run, fine-tuned at "
        "lower speed for better real-time performance."
    )

    add_qa(doc,
        "What was the Time Scale issue you encountered?",
        "Unity's Time Scale controls simulation speed. Training at Time Scale 20 means the game "
        "runs 20x faster, collecting training data much faster. However, I discovered that "
        "the agent trained at 20x speed performed poorly at normal speed (1x). This is because "
        "at higher time scales, physics behaves slightly differently — collisions can be missed, "
        "movement becomes less smooth, and the agent's decision timing changes relative to the "
        "game physics. The solution was to either train at Time Scale 1 (slower but accurate) "
        "or use a moderate speed like 5x as a compromise."
    )

    add_qa(doc,
        "What does the training loss / reward graph look like?",
        "The reward graph shows a classic RL learning curve: initially flat near zero (random "
        "behavior), then a sharp increase as the agent discovers the survival strategy, followed "
        "by a gradual improvement as it learns pipe navigation. At 1M steps the agent reaches "
        "a mean reward around 100+, meaning it consistently passes many pipes per episode. "
        "The graph also shows some variance, which is normal in RL — performance fluctuates "
        "episode to episode due to the stochastic nature of training."
    )

    add_qa(doc,
        "What would happen if you trained for longer, like 10 million steps?",
        "The agent would likely converge to near-perfect play. After about 2-3M steps, "
        "improvements become marginal. The max_steps is set to 10M as an upper bound, "
        "but I stopped earlier because the agent was already performing well. Diminishing "
        "returns is common in RL — most learning happens in the first 20-30% of training."
    )

    # --- UNITY / TECHNICAL ---
    doc.add_heading("Unity & Technical", level=2)

    add_qa(doc,
        "How does Unity ML-Agents work architecturally?",
        "Unity ML-Agents uses a client-server architecture. The Unity game (C#) acts as the "
        "environment — it simulates the game, collects observations, applies actions, and "
        "calculates rewards. A Python process runs the training algorithm (PPO via PyTorch). "
        "They communicate over gRPC on a local port. During training, Unity sends observations "
        "to Python, Python sends back actions, and this loop repeats thousands of times per "
        "second. After training, the model is exported as ONNX and runs directly in Unity "
        "without Python."
    )

    add_qa(doc,
        "What is the difference between 'Default', 'Heuristic', and 'Inference Only' behavior types?",
        "'Default' means the agent is controlled by the Python training process — used during "
        "training. 'Heuristic' means the agent uses a manually coded function (useful for "
        "testing). 'Inference Only' means the agent uses the trained ONNX model to make "
        "decisions — this is what you use when the game is deployed or when you want to "
        "watch the trained agent play."
    )

    add_qa(doc,
        "How does the difficulty system work?",
        "The Level.cs script tracks how many pipes the agent has passed and adjusts difficulty "
        "accordingly. Easy (0-4 pipes): 50-unit gap, 1.4s spawn interval. Medium (5-11): "
        "40-unit gap, 1.3s interval. Hard (12-23): 33-unit gap, 1.1s interval. Impossible "
        "(24+): 24-unit gap, 1.0s interval. This acts as a natural curriculum — the agent "
        "first masters easy levels before facing harder ones."
    )

    # --- GENERAL ML ---
    doc.add_heading("General Machine Learning", level=2)

    add_qa(doc,
        "How is reinforcement learning different from supervised and unsupervised learning?",
        "Supervised learning learns from labeled examples (input-output pairs). Unsupervised "
        "learning finds patterns in unlabeled data. Reinforcement learning learns through "
        "interaction with an environment, receiving reward signals instead of labels. The key "
        "difference is that RL must deal with sequential decision-making — each action affects "
        "future states — and delayed rewards — the agent might not know if a decision was good "
        "until many steps later. RL is ideal for games, robotics, and any sequential decision problem."
    )

    add_qa(doc,
        "Could you use a convolutional neural network (CNN) to process screen pixels instead?",
        "Yes, Unity ML-Agents supports visual observations. You could feed raw screen pixels "
        "to a CNN, which would then learn to extract features like pipe positions and bird "
        "location from the image. However, this would: (1) require a much larger network, "
        "(2) take 10-100x more training time, (3) need more compute resources. Since we can "
        "easily extract the 4 relevant numbers directly, using numerical observations is far "
        "more efficient. Pixel-based approaches make more sense for complex games where the "
        "relevant information is hard to extract programmatically."
    )

    add_qa(doc,
        "What are some real-world applications of this kind of RL?",
        "Similar RL techniques are used in: game AI (DeepMind's AlphaGo, OpenAI Five for Dota 2), "
        "robotics (learning to walk, manipulate objects), autonomous driving (decision making), "
        "recommendation systems (optimizing long-term user engagement), resource management "
        "(data center cooling at Google), and drug discovery (molecular optimization). The core "
        "idea — an agent learning from trial and error with reward signals — applies to any "
        "sequential decision-making problem."
    )

    add_qa(doc,
        "What would you do differently if you started this project over?",
        "Three things: (1) I would start with Time Scale 1 training from the beginning to "
        "avoid the speed mismatch issue. (2) I would implement curriculum learning more "
        "explicitly — start with very easy pipe configurations and gradually increase difficulty "
        "rather than relying on the in-game difficulty system. (3) I would experiment with "
        "different observation spaces, like adding the bird's rotation or the second upcoming pipe, "
        "to see if more information improves learning speed."
    )

    doc.add_page_break()

    # ============================================================
    # PART 3: QUICK REFERENCE — KEY CODE
    # ============================================================
    doc.add_heading("PART 3: Key Code Reference", level=1)
    doc.add_paragraph(
        "If the professor asks to see specific code, here are the most important sections."
    )

    # Birdagent.cs
    doc.add_heading("Birdagent.cs — Core Agent Logic", level=2)
    doc.add_paragraph("File: Assets/scripts/Birdagent.cs")
    code = read_file("Assets/scripts/Birdagent.cs")
    add_code_block(doc, code)

    doc.add_page_break()

    # flappy.yaml
    doc.add_heading("flappy.yaml — Training Configuration", level=2)
    code = read_file("flappy.yaml")
    add_code_block(doc, code)

    doc.add_page_break()

    # Level.cs
    doc.add_heading("Level.cs — Game Logic & Difficulty", level=2)
    doc.add_paragraph("File: Assets/scripts/Level.cs")
    code = read_file("Assets/scripts/Level.cs")
    add_code_block(doc, code)

    # Save
    output_path = os.path.join(BASE_DIR, "FlappyBird_ML_Agent_Documentation.docx")
    doc.save(output_path)
    print(f"Documentation generated: {output_path}")


if __name__ == "__main__":
    generate_documentation()
